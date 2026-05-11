"""
DOCX Translation with Azure OpenAI and Formatting Preservation

This script reads a .docx file, translates the text content using Azure OpenAI
while preserving all formatting, fonts, and layout, then outputs a new translated .docx file.

"""
import os
import re
import time
from typing import List
from docx import Document
from openai import AzureOpenAI
# local imports
import utils


def translate_text_batch(client: AzureOpenAI, deployment_name: str, texts: List[str], target_language: str) -> List[str]:
    """
    Translate multiple text strings in a single API call for efficiency.

    Args:
        client: Azure OpenAI client
        deployment_name: Azure OpenAI deployment name
        texts: List of texts to translate
        target_language: Target language

    Returns:
        List of translated texts
    """
    if not texts or all(not text.strip() for text in texts):
        return texts

    # Filter out empty texts but keep track of their positions
    text_map = {}
    non_empty_texts = []
    for i, text in enumerate(texts):
        if text.strip():
            text_map[len(non_empty_texts)] = i
            non_empty_texts.append(text)

    if not non_empty_texts:
        return texts

    prompt = f"""Translate each <text> element below to {target_language}.
Preserve the exact meaning, tone, and any special characters.
Preserve any inline <i>...</i> and <b>...</b> markup tags within the text content.
Return ONLY the translated text for each element using the identical tag structure:
<text id="1">translation here</text>
<text id="2">translation here</text>
Do not add any text outside the tags.

Texts to translate:
"""
    for i, text in enumerate(non_empty_texts):
        prompt += f'\n<text id="{i+1}">{text}</text>'

    try:
        response = client.chat.completions.create(
            model=deployment_name,
            messages=[
                {"role": "system", "content": "You are a professional translator. Translate accurately while preserving formatting and meaning. Keep any inline <i>...</i> and <b>...</b> markup tags in place around the translated words."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=4000,
            temperature=0
        )

        translated_content = response.choices[0].message.content.strip()

        # Parse <text id="N">...</text> tags
        matches = re.findall(r'<text id="(\d+)">(.*?)</text>', translated_content, re.DOTALL)

        result = texts.copy()
        if matches:
            for id_str, translation in matches:
                idx = int(id_str) - 1
                if 0 <= idx < len(non_empty_texts):
                    original_index = text_map[idx]
                    result[original_index] = translation.strip()
        else:
            # Last-resort fallback: old separator approach
            if "---TRANSLATION_SEPARATOR---" in translated_content:
                translated_parts = translated_content.split("---TRANSLATION_SEPARATOR---")
                for i, translation in enumerate(translated_parts[:len(non_empty_texts)]):
                    original_index = text_map[i]
                    result[original_index] = translation.strip()

        return result

    except Exception as e:
        print(f"Translation error: {e}")
        # Return original texts if translation fails
        return texts


def _parse_inline_markup(text: str) -> List[tuple]:
    """Split text containing <i>/<b> tags into (segment_text, fmt_dict) pairs."""
    parts = re.split(r'(<[ib]>.*?</[ib]>)', text, flags=re.DOTALL)
    segments = []
    for part in parts:
        if not part:
            continue
        m = re.match(r'<(i|b)>(.*?)</\1>', part, re.DOTALL)
        if m:
            tag, content = m.group(1), m.group(2)
            segments.append((content, {'italic': True} if tag == 'i' else {'bold': True}))
        else:
            segments.append((part, {}))
    return segments


def _apply_run_formatting(run, base_fmt: dict, seg_fmt: dict):
    """Apply base paragraph formatting plus per-segment overrides to a run."""
    if base_fmt:
        if base_fmt['font_name']:
            run.font.name = base_fmt['font_name']
        if base_fmt['font_size']:
            run.font.size = base_fmt['font_size']
        if base_fmt['font_color']:
            run.font.color.rgb = base_fmt['font_color']
        if base_fmt['highlight_color']:
            run.font.highlight_color = base_fmt['highlight_color']
        if base_fmt['style']:
            run.style = base_fmt['style']
        bold = seg_fmt.get('bold', base_fmt['bold'])
        italic = seg_fmt.get('italic', base_fmt['italic'])
        underline = seg_fmt.get('underline', base_fmt['underline'])
    else:
        bold = seg_fmt.get('bold')
        italic = seg_fmt.get('italic')
        underline = seg_fmt.get('underline')
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline is not None:
        run.underline = underline


def collect_paragraph_texts(paragraphs) -> List[str]:
    """
    Collect text from paragraphs, encoding italic/bold runs as <i>/<b> inline tags.

    Returns:
        List of markup-annotated text strings
    """
    texts = []
    for paragraph in paragraphs:
        if not paragraph.text.strip():
            texts.append("")
            continue
        marked = ""
        for run in paragraph.runs:
            t = run.text
            if not t:
                continue
            if run.italic:
                t = f"<i>{t}</i>"
            elif run.bold:
                t = f"<b>{t}</b>"
            marked += t
        texts.append(marked if marked else paragraph.text)
    return texts


def apply_translations_to_paragraphs(paragraphs, translations: List[str]):
    """
    Apply translations to paragraphs, reconstructing per-run italic/bold from inline markup.

    Args:
        paragraphs: List of paragraph objects
        translations: List of translated texts (may contain <i>/<b> inline markup)
    """
    for paragraph, translation in zip(paragraphs, translations):
        if not translation.strip():
            continue

        # Capture base formatting from first meaningful run
        base_fmt = None
        for run in paragraph.runs:
            if run.text.strip():
                base_fmt = {
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name,
                    'font_size': run.font.size,
                    'font_color': run.font.color.rgb if run.font.color.rgb else None,
                    'highlight_color': run.font.highlight_color,
                    'style': run.style
                }
                break

        # Clear existing runs and keep just one
        for run in paragraph.runs:
            run.clear()
        while len(paragraph.runs) > 1:
            paragraph._element.remove(paragraph.runs[-1]._element)

        # Parse inline markup and rebuild runs
        segments = _parse_inline_markup(translation) or [(translation, {})]

        if paragraph.runs:
            first_text, first_fmt = segments[0]
            paragraph.runs[0].text = first_text
            _apply_run_formatting(paragraph.runs[0], base_fmt, first_fmt)
            for text, seg_fmt in segments[1:]:
                _apply_run_formatting(paragraph.add_run(text), base_fmt, seg_fmt)


def translate_table_cells(client: AzureOpenAI, model: str, table, target_language: str):
    """
    Translate all text in table cells using batch processing.

    Args:
        table: python-docx table object
    """
    all_paragraphs = []

    # Collect all paragraphs from all cells
    for row in table.rows:
        for cell in row.cells:
            all_paragraphs.extend(cell.paragraphs)

    if not all_paragraphs:
        return

    # Collect texts and translate in batches
    texts = collect_paragraph_texts(all_paragraphs)

    # Process in batches of 20 to avoid token limits
    batch_size = 20
    translated_texts = []

    for i in range(0, len(texts), batch_size):
        batch = texts[i:i + batch_size]
        batch_translations = translate_text_batch(client=client,
                                                  deployment_name=model,
                                                  texts=batch,
                                                  target_language=target_language)
        translated_texts.extend(batch_translations)

        # Small delay between batches to avoid rate limiting
        if i + batch_size < len(texts):
            time.sleep(0.5)

    # Apply translations
    apply_translations_to_paragraphs(all_paragraphs, translated_texts)


def translate_docx_document(client: AzureOpenAI,
                            model: str,
                            input_path: str,
                            target_language: str,
                            output_folder: str,
                            output_format: str) -> bool:
    """
    Translate an entire DOCX document while preserving formatting.

    Args:
        client: Azure OpenAI client
        model: chosen Azure OpenAI model deployment
        input_path: path to input .docx file
        target_language: language to translate to
        output_folder: output folder name
        output_format: chosen output format

    Returns:
        True if successful, False otherwise
    """
    try:
        file_name = os.path.basename(input_path)
        output_file_path = os.path.join(output_folder, target_language + "_" + file_name)
        # Load the document
        doc = Document(input_path)

        # Translate main document paragraphs
        if doc.paragraphs:
            texts = collect_paragraph_texts(doc.paragraphs)

            # Process in batches
            batch_size = 20
            translated_texts = []

            for i in range(0, len(texts), batch_size):
                batch = texts[i:i + batch_size]
                print(f"Processing batch {i//batch_size + 1}/{(len(texts)-1)//batch_size + 1}")
                batch_translations = translate_text_batch(client=client,
                                                          deployment_name=model,
                                                          texts=batch,
                                                          target_language=target_language)
                translated_texts.extend(batch_translations)

                # Small delay between batches
                if i + batch_size < len(texts):
                    time.sleep(0.5)

            apply_translations_to_paragraphs(doc.paragraphs, translated_texts)

        # Translate tables
        if doc.tables:
            print("Translating tables...")
            for table_idx, table in enumerate(doc.tables):
                print(f"Processing table {table_idx+1}/{len(doc.tables)}")
                translate_table_cells(client=client,
                                      model=model,
                                      table=table,
                                      target_language=target_language)

        # Translate headers and footers
        print("Translating headers and footers...")
        for section_idx, section in enumerate(doc.sections):
            print(f"Processing section {section_idx+1}/{len(doc.sections)}")

            # Translate header
            if section.header and section.header.paragraphs:
                header_texts = collect_paragraph_texts(section.header.paragraphs)
                header_translations = translate_text_batch(client=client,
                                                           deployment_name=model,
                                                           texts=header_texts,
                                                           target_language=target_language)
                apply_translations_to_paragraphs(section.header.paragraphs, header_translations)

                # Translate header tables
                for table in section.header.tables:
                    translate_table_cells(client=client,
                                          model=model,
                                          table=table,
                                          target_language=target_language)

            # Translate footer
            if section.footer and section.footer.paragraphs:
                footer_texts = collect_paragraph_texts(section.footer.paragraphs)
                footer_translations = translate_text_batch(client=client,
                                                           deployment_name=model,
                                                           texts=footer_texts,
                                                           target_language=target_language)
                apply_translations_to_paragraphs(section.footer.paragraphs, footer_translations)

                # Translate footer tables
                for table in section.footer.tables:
                    translate_table_cells(client=client,
                                          model=model,
                                          table=table,
                                          target_language=target_language)

        # if indicated, save as pdf file
        if output_format == "Save as PDF":
            pdf_file_name = os.path.splitext(file_name)[0] + ".pdf"
            pdf_file_path = os.path.join(output_folder, target_language + "_" + pdf_file_name)
            utils.convert_docx_to_pdf(output_file_path, pdf_file_path)
            # add watermark to created pdf file
            watermark_file_path = os.path.abspath(os.path.join(os.getcwd(), "watermark.pdf"))
            utils.add_watermark(pdf_file_path, pdf_file_path, watermark_file_path)
            # remove converted .docx file
            os.remove(output_file_path)
        elif output_format == "Save as plain text":
            # Convert to plain text file
            txt_file_name = os.path.splitext(file_name)[0] + ".txt"
            txt_file_path = os.path.join(output_folder, target_language + "_" + txt_file_name)
            # utils.convert_docx_to_txt(output_file_path, txt_file_path)
            # remove converted .docx file
            os.remove(output_file_path)
        else:
            # Save the translated document
            print(f"Saving translated document: {output_file_path}")
            doc.save(output_file_path)

        return True

    except Exception as e:
        print(f"Error processing document: {e}")
        import traceback
        traceback.print_exc()
        return False
